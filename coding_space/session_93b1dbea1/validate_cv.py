#!/usr/bin/env python3
"""Validate CV files generated"""
import os
from docx import Document

print('=== File Existence & Sizes ===')
docx_path = 'CV_Tim_Chu.docx'
pdf_path = 'CV_Tim_Chu.pdf'

for f in [docx_path, pdf_path]:
    if os.path.exists(f):
        size = os.path.getsize(f)
        print(f'  [OK] {f}: {size} bytes')
    else:
        print(f'  [ERROR] {f} not found!')
        exit(1)

print()
print('=== DOCX Content Verification ===')
doc = Document(docx_path)
print(f'Total paragraphs: {len(doc.paragraphs)}')

full_text = ' '.join([p.text for p in doc.paragraphs])

checks = [
    'Tim Chu',
    'Technical Business Analyst',
    'NESA',
    '14+ portals',
    'CVSS',
    'penetration testing',
    'malicious file scanning',
    '500+ member',
    'Cyber Security Analyst',
    '7+ years',
]

all_ok = True
for check in checks:
    if check.lower() in full_text.lower():
        print(f'  [OK] Found: {check}')
    else:
        print(f'  [MISSING] {check}')
        all_ok = False

print(f'\nAll checks passed: {all_ok}')

# Also verify PDF has reasonable size
if os.path.getsize(pdf_path) > 1000:
    print('[OK] PDF has reasonable content size')
else:
    print('[WARNING] PDF may be too small')

if all_ok:
    print('\n=== VALIDATION SUCCESSFUL ===')
    exit(0)
else:
    print('\n=== VALIDATION FAILED ===')
    exit(1)