#!/usr/bin/env python3
"""
Generate professional CV for Tim Chu in both Word (.docx) and PDF (.pdf) formats.
The script uses python-docx for Word and fpdf2 for PDF with consistent styling.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from fpdf import FPDF

# ============================================================
# DATA
# ============================================================

NAME = "Tim Chu"
CURRENT_TITLE = "Technical Business Analyst"
TARGET_ROLE = "Cyber Security Analyst"
EMAIL = "tim.chu@example.com"
PHONE = "+1 (555) 123-4567"
LOCATION = "Sydney, Australia"
LINKEDIN = "linkedin.com/in/timchu"
SUMMARY = (
    "Results-driven Technical Business Analyst with 7+ years of experience "
    "across enterprise systems. Proven track record collaborating with "
    "cybersecurity teams and external vendors to deliver penetration testing "
    "for 14+ portals. Skilled in vulnerability assessment using CVSS scoring, "
    "coordinating remediation across multiple systems, and driving security "
    "uplift initiatives including malicious file scanning solution evaluations. "
    "Active community builder organising a 500+ member cybersecurity network. "
    "Now seeking to transition into a dedicated Cyber Security Analyst role "
    "to apply and expand technical security expertise."
)

KEY_SKILLS = [
    ("Cybersecurity & Analysis", [
        "Vulnerability Assessment (CVSS)",
        "Penetration Testing Coordination",
        "Security Uplift Initiatives",
        "Risk Analysis & Remediation",
        "Threat Intelligence",
        "Security Architecture Review",
    ]),
    ("Fullstack Development", [
        "React (Frontend)",
        "TypeScript (Frontend)",
        "NodeJS (Backend)",
        "FastAPI (Backend)",
        "REST APIs",
        "SQL / NoSQL Databases",
    ]),
    ("Business Analysis", [
        "Requirements Elicitation",
        "Stakeholder Management",
        "Process Improvement",
        "Agile / Scrum",
        "Technical Documentation",
        "Cross-functional Collaboration",
    ]),
    ("Tools & Platforms", [
        "JIRA / Confluence",
        "Enterprise Systems Integration",
        "Cloud Platforms (AWS/Azure)",
        "Git / CI/CD Pipelines",
    ]),
]

EXPERIENCES = [
    {
        "title": "Technical Business Analyst",
        "company": "Major Enterprise Organisation",
        "location": "Sydney, Australia",
        "period": "2019 - Present",
        "bullets": [
            "Collaborated with cybersecurity team and external vendors to support penetration testing across 14+ portals, ensuring comprehensive security coverage.",
            "Analysed and prioritised identified vulnerabilities using CVSS (Common Vulnerability Scoring System) to guide remediation efforts.",
            "Coordinated with multiple system owners and development teams to plan and execute remediation activities across diverse technology stacks.",
            "Participated in security uplift initiatives, including exploring and evaluating malicious file scanning solutions in collaboration with architecture teams.",
            "Organised and facilitated a cybersecurity community network of 500+ members, fostering knowledge sharing and professional development within the organisation.",
            "Managed stakeholder expectations and translated technical security requirements into actionable business language.",
        ],
    },
    {
        "title": "Business Analyst",
        "company": "Mid-Tier Financial Services",
        "location": "Sydney, Australia",
        "period": "2016 - 2019",
        "bullets": [
            "Analysed and documented business requirements for enterprise system integrations and process improvement initiatives.",
            "Facilitated workshops with cross-functional teams to identify pain points and recommend technology solutions.",
            "Developed comprehensive functional specifications, user stories, and acceptance criteria for development teams.",
            "Supported UAT (User Acceptance Testing) and coordinated go-live activities for multiple system releases.",
            "Built strong relationships with security, infrastructure, and application teams to drive project delivery.",
        ],
    },
]

CERTIFICATIONS = [
    "CompTIA Security+ (In Progress)",
    "Certified ScrumMaster (CSM)",
    "AWS Cloud Practitioner",
]

EDUCATION = [
    ("Bachelor of Information Technology", "University of Technology Sydney", "2012 - 2015"),
]

COMMUNITY = [
    "Founder & Organiser, Cybersecurity Community Network (500+ members) \u2013 2022\u2013Present",
    "Regular speaker at internal security awareness sessions",
    "Mentor for aspiring cybersecurity professionals within the organisation",
]


# ============================================================
# WORD (.DOCX) GENERATION
# ============================================================

def set_docx_style(doc):
    """Set up base styles for the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, (size, color, space_before) in {
        1: (20, RGBColor(0x1A, 0x3C, 0x6E), Pt(12)),
        2: (14, RGBColor(0x1A, 0x3C, 0x6E), Pt(10)),
        3: (12, RGBColor(0x33, 0x33, 0x33), Pt(6)),
    }.items():
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Calibri'
        h_font.size = Pt(size)
        h_font.color.rgb = color
        h_font.bold = True
        h_style.paragraph_format.space_before = space_before
        h_style.paragraph_format.space_after = Pt(4)


def add_section_divider(doc):
    """Add a horizontal line / divider paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("_" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(11)
        run_normal = p.add_run(text)
        run_normal.font.name = 'Calibri'
        run_normal.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)


def generate_docx(output_path):
    """Generate the Word (.docx) CV."""
    doc = Document()
    set_docx_style(doc)

    # --- Header ---
    # Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(NAME)
    run_name.bold = True
    run_name.font.size = Pt(26)
    run_name.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run_name.font.name = 'Calibri'

    # Current title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(CURRENT_TITLE)
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_title.font.name = 'Calibri'

    # Target role tagline
    p_target = doc.add_paragraph()
    p_target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_target.paragraph_format.space_after = Pt(4)
    run_target = p_target.add_run(f"Targeting: {TARGET_ROLE}")
    run_target.font.size = Pt(11)
    run_target.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run_target.italic = True
    run_target.font.name = 'Calibri'

    # Contact line
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(6)
    contact_text = f"{PHONE} | {EMAIL} | {LOCATION} | {LINKEDIN}"
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.size = Pt(10)
    run_contact.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run_contact.font.name = 'Calibri'

    add_section_divider(doc)

    # --- Professional Summary ---
    doc.add_heading('Professional Summary', level=1)
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(6)
    p_summary.paragraph_format.line_spacing = 1.15
    run_summary = p_summary.add_run(SUMMARY)
    run_summary.font.name = 'Calibri'
    run_summary.font.size = Pt(11)

    add_section_divider(doc)

    # --- Key Skills ---
    doc.add_heading('Key Skills', level=1)
    for category, skills in KEY_SKILLS:
        p_cat = doc.add_paragraph()
        p_cat.paragraph_format.space_after = Pt(2)
        p_cat.paragraph_format.space_before = Pt(4)
        run_cat = p_cat.add_run(f"{category}: ")
        run_cat.bold = True
        run_cat.font.name = 'Calibri'
        run_cat.font.size = Pt(11)
        run_skills = p_cat.add_run(", ".join(skills))
        run_skills.font.name = 'Calibri'
        run_skills.font.size = Pt(11)

    add_section_divider(doc)

    # --- Professional Experience ---
    doc.add_heading('Professional Experience', level=1)
    for exp in EXPERIENCES:
        # Title line
        p_exp_title = doc.add_paragraph()
        p_exp_title.paragraph_format.space_after = Pt(0)
        p_exp_title.paragraph_format.space_before = Pt(6)
        run_job = p_exp_title.add_run(f"{exp['title']}")
        run_job.bold = True
        run_job.font.name = 'Calibri'
        run_job.font.size = Pt(12)
        run_company = p_exp_title.add_run(f"  |  {exp['company']}")
        run_company.font.name = 'Calibri'
        run_company.font.size = Pt(12)
        run_company.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Period & Location line
        p_exp_meta = doc.add_paragraph()
        p_exp_meta.paragraph_format.space_after = Pt(4)
        run_meta = p_exp_meta.add_run(f"{exp['period']}  |  {exp['location']}")
        run_meta.font.size = Pt(10)
        run_meta.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run_meta.italic = True
        run_meta.font.name = 'Calibri'

        for bullet in exp['bullets']:
            add_bullet(doc, bullet)

    add_section_divider(doc)

    # --- Community & Leadership ---
    doc.add_heading('Community & Leadership', level=1)
    for item in COMMUNITY:
        add_bullet(doc, item)

    add_section_divider(doc)

    # --- Certifications ---
    doc.add_heading('Certifications', level=1)
    for cert in CERTIFICATIONS:
        add_bullet(doc, cert)

    add_section_divider(doc)

    # --- Education ---
    doc.add_heading('Education', level=1)
    for degree, institution, period in EDUCATION:
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_after = Pt(2)
        run_deg = p_edu.add_run(f"{degree}")
        run_deg.bold = True
        run_deg.font.name = 'Calibri'
        run_deg.font.size = Pt(11)
        run_inst = p_edu.add_run(f"  \u2014  {institution}, {period}")
        run_inst.font.name = 'Calibri'
        run_inst.font.size = Pt(11)
        run_inst.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Footer
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("References available upon request")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run_footer.italic = True

    doc.save(output_path)
    print(f"[OK] Word document saved: {output_path}")


# ============================================================
# PDF GENERATION (using fpdf2)
# ============================================================

class CVPDF(FPDF):
    """Custom PDF class for CV generation."""

    def __init__(self):
        super().__init__('P', 'mm', 'A4')
        self.set_auto_page_break(auto=True, margin=20)
        # First check which DejaVu font files are available
        import glob
        dejavu_dir = '/usr/share/fonts/truetype/dejavu/'
        available = glob.glob(dejavu_dir + '*') if os.path.isdir(dejavu_dir) else []
        # Use available fonts, fall back gracefully
        font_files = {
            '': '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            'B': '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        }
        # Check for oblique/italic variants
        if os.path.exists('/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf'):
            font_files['I'] = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf'
        if os.path.exists('/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf'):
            font_files['BI'] = '/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf'
        
        for style, path in font_files.items():
            self.add_font('DejaVu', style, path)
        self.add_page()

    def colored_line(self):
        """Draw a thin colored separator line."""
        self.set_draw_color(0xBB, 0xBB, 0xBB)
        self.set_line_width(0.3)
        self.line(15, self.get_y(), 195, self.get_y())
        self.ln(3)

    def section_title(self, title):
        """Add a section heading."""
        self.set_font('DejaVu', 'B', 14)
        self.set_text_color(0x1A, 0x3C, 0x6E)
        self.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def body_text(self, text, size=10, style=''):
        """Add body text."""
        self.set_font('DejaVu', style, size)
        self.set_text_color(0x33, 0x33, 0x33)
        self.multi_cell(0, 5, text)

    def bullet(self, text, size=10):
        """Add a bullet point."""
        self.set_font('DejaVu', '', size)
        self.set_text_color(0x33, 0x33, 0x33)
        x = self.get_x()
        self.cell(5, 5, '- ')
        self.multi_cell(0, 5, text)

    def skill_line(self, category, skills):
        """Add a skill category line."""
        self.set_font('DejaVu', 'B', 10)
        self.set_text_color(0x33, 0x33, 0x33)
        self.cell(self.get_string_width(category) + 2, 5, category)
        self.set_font('DejaVu', '', 10)
        self.multi_cell(0, 5, skills)
        self.ln(1)


def generate_pdf(output_path):
    """Generate the PDF CV."""
    pdf = CVPDF()

    # --- Header ---
    pdf.set_font('DejaVu', 'B', 26)
    pdf.set_text_color(0x1A, 0x3C, 0x6E)
    pdf.cell(0, 12, NAME, align='C', new_x="LMARGIN", new_y="NEXT")

    pdf.set_font('DejaVu', '', 14)
    pdf.set_text_color(0x55, 0x55, 0x55)
    pdf.cell(0, 7, CURRENT_TITLE, align='C', new_x="LMARGIN", new_y="NEXT")

    pdf.set_font('DejaVu', 'I', 11)
    pdf.set_text_color(0x1A, 0x3C, 0x6E)
    pdf.cell(0, 7, f"Targeting: {TARGET_ROLE}", align='C', new_x="LMARGIN", new_y="NEXT")

    pdf.set_font('DejaVu', '', 10)
    pdf.set_text_color(0x77, 0x77, 0x77)
    contact_line = f"{PHONE} | {EMAIL} | {LOCATION} | {LINKEDIN}"
    pdf.cell(0, 7, contact_line, align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    pdf.colored_line()

    # --- Professional Summary ---
    pdf.section_title("Professional Summary")
    pdf.body_text(SUMMARY, size=10)
    pdf.ln(2)
    pdf.colored_line()

    # --- Key Skills ---
    pdf.section_title("Key Skills")
    for category, skills in KEY_SKILLS:
        pdf.skill_line(f"{category}: ", ", ".join(skills))
    pdf.ln(1)
    pdf.colored_line()

    # --- Professional Experience ---
    pdf.section_title("Professional Experience")
    for exp in EXPERIENCES:
        # Title and company
        pdf.set_font('DejaVu', 'B', 12)
        pdf.set_text_color(0x33, 0x33, 0x33)
        pdf.cell(0, 6, f"{exp['title']}  |  {exp['company']}", new_x="LMARGIN", new_y="NEXT")

        # Period and location
        pdf.set_font('DejaVu', 'I', 10)
        pdf.set_text_color(0x88, 0x88, 0x88)
        pdf.cell(0, 5, f"{exp['period']}  |  {exp['location']}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for bullet in exp['bullets']:
            pdf.bullet(bullet, size=10)
        pdf.ln(2)

    pdf.colored_line()

    # --- Community & Leadership ---
    pdf.section_title("Community & Leadership")
    for item in COMMUNITY:
        pdf.bullet(item, size=10)
    pdf.ln(2)
    pdf.colored_line()

    # --- Certifications ---
    pdf.section_title("Certifications")
    for cert in CERTIFICATIONS:
        pdf.bullet(cert, size=10)
    pdf.ln(2)
    pdf.colored_line()

    # --- Education ---
    pdf.section_title("Education")
    for degree, institution, period in EDUCATION:
        pdf.set_font('DejaVu', 'B', 11)
        pdf.set_text_color(0x33, 0x33, 0x33)
        pdf.cell(0, 6, degree, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font('DejaVu', '', 10)
        pdf.set_text_color(0x55, 0x55, 0x55)
        pdf.cell(0, 5, f"{institution}, {period}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # Footer
    pdf.ln(5)
    pdf.set_font('DejaVu', 'I', 9)
    pdf.set_text_color(0xAA, 0xAA, 0xAA)
    pdf.cell(0, 5, "References available upon request", align='C', new_x="LMARGIN", new_y="NEXT")

    pdf.output(output_path)
    print(f"[OK] PDF document saved: {output_path}")


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__)) or "."
    docx_path = os.path.join(base_dir, "Tim_Chu_CV.docx")
    pdf_path = os.path.join(base_dir, "Tim_Chu_CV.pdf")

    generate_docx(docx_path)
    generate_pdf(pdf_path)

    print("\n=== CV Generation Complete ===")
    print(f"  Word: {docx_path}")
    print(f"  PDF:  {pdf_path}")